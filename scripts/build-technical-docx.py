#!/usr/bin/env python3
"""Build the styled MeetBroker technical DOCX from the canonical Markdown."""

from __future__ import annotations

import re
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "TECHNICAL_OVERVIEW.md"
OUTPUT = ROOT / "docs" / "MeetBroker-Technical-Overview.docx"
BRAND_MARK = ROOT / "docs" / "screenshots" / "brand-mark.png"

NAVY = "071A3A"
NAVY_2 = "0B2B5A"
COBALT = "1488FF"
CYAN = "51D7FF"
PINK = "FF4E91"
ICE = "EAF4FF"
ICE_2 = "D9ECFF"
TEXT = "102446"
MUTED = "5E7292"
WHITE = "FFFFFF"
LINE = "A9CDED"


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_paragraph_bottom_border(paragraph, color=CYAN, size="14") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_repeatable_heading(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("MEETBROKER / TECHNICAL OVERVIEW   ·   ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        style.font.size = Pt(10)
        style.font.color.rgb = RGBColor.from_string(TEXT)

    for style_name, size, color in (
        ("Title", 30, NAVY),
        ("Heading 1", 19, NAVY),
        ("Heading 2", 14, NAVY_2),
        ("Heading 3", 11, COBALT),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.keep_with_next = True

    header = section.header
    table = header.add_table(1, 2, Inches(6.95))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.85)
    table.columns[1].width = Inches(2.1)
    set_table_borders(table, NAVY, "0")
    for cell in table.rows[0].cells:
        set_cell_fill(cell, NAVY)
        set_cell_margins(cell, 80, 120, 80, 120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    if BRAND_MARK.exists():
        left.add_run().add_picture(str(BRAND_MARK), height=Cm(0.55))
    brand = left.add_run("  MEETBROKER")
    brand.font.name = "Arial"
    brand.font.bold = True
    brand.font.size = Pt(10)
    brand.font.color.rgb = RGBColor.from_string(WHITE)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    marker = right.add_run("SYSTEM / 2026")
    marker.font.name = "Arial"
    marker.font.bold = True
    marker.font.size = Pt(8)
    marker.font.color.rgb = RGBColor.from_string(CYAN)
    set_paragraph_bottom_border(header.add_paragraph(), CYAN, "20")

    add_page_number(section.footer.paragraphs[0])


def add_inline(paragraph, text: str, *, color: str | None = None) -> None:
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(COBALT)
        else:
            run = paragraph.add_run(part)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def force_run_font(run, name: str) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)


def add_picture(run, image_path: Path, *, width=None, height=None) -> None:
    if image_path.suffix.lower() != ".webp":
        run.add_picture(str(image_path), width=width, height=height)
        return
    with tempfile.NamedTemporaryFile(suffix=".png") as converted:
        with Image.open(image_path) as source:
            source.convert("RGB").save(converted.name, format="PNG", optimize=True)
        run.add_picture(converted.name, width=width, height=height)


def add_cover(document: Document) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)

    label = document.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = label.add_run("CORPORATE MEETING INFRASTRUCTURE  /  01")
    run.font.name = "Arial"
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(COBALT)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run("MEET")
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(35)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    run = title.add_run("BROKER")
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(35)
    run.font.color.rgb = RGBColor.from_string(COBALT)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("ТЕХНІЧНИЙ ОПИС СИСТЕМИ")
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    set_paragraph_bottom_border(subtitle, PINK, "18")

    summary = document.add_paragraph()
    summary.paragraph_format.space_before = Pt(10)
    summary.paragraph_format.space_after = Pt(14)
    add_inline(
        summary,
        "Архітектура · модель даних · безпека · сповіщення · "
        "контейнеризація · CI/CD",
        color=MUTED,
    )

    if BRAND_MARK.exists():
        mark_line = document.add_paragraph()
        mark_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_picture(mark_line.add_run(), BRAND_MARK, height=Cm(1.7))

    hero = ROOT / "docs" / "screenshots" / "my-meetings-light.webp"
    if hero.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(8)
        add_picture(paragraph.add_run(), hero, width=Inches(6.75))

    meta = document.add_table(1, 3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = True
    values = (
        ("СТАН", "Конкурсне MVP"),
        ("ВЕРСІЯ", "1.0"),
        ("ДАТА", "31.07.2026"),
    )
    for index, (label_text, value) in enumerate(values):
        cell = meta.cell(0, index)
        set_cell_fill(cell, NAVY if index != 1 else NAVY_2)
        set_cell_margins(cell, 110, 130, 110, 130)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        label_run = p.add_run(f"{label_text}\n")
        label_run.bold = True
        label_run.font.size = Pt(7)
        label_run.font.color.rgb = RGBColor.from_string(CYAN)
        value_run = p.add_run(value)
        value_run.bold = True
        value_run.font.size = Pt(9)
        value_run.font.color.rgb = RGBColor.from_string(WHITE)
    set_table_borders(meta, WHITE, "0")


def add_contents(document: Document, lines: list[str]) -> None:
    heading = document.add_paragraph()
    run = heading.add_run("ЗМІСТ")
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    set_paragraph_bottom_border(heading, CYAN, "18")

    intro = document.add_paragraph(
        "Документ сформовано з канонічного Markdown у репозиторії. "
        "Нумерація відповідає розділам технічного опису."
    )
    intro.paragraph_format.space_after = Pt(12)

    table = document.add_table(0, 2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(1.25)
    table.columns[1].width = Cm(14.8)
    for title in lines:
        match = re.match(r"(\d+)\.\s+(.+)", title)
        number = match.group(1) if match else "·"
        name = match.group(2) if match else title
        cells = table.add_row().cells
        set_cell_fill(cells[0], COBALT)
        set_cell_fill(cells[1], ICE)
        set_cell_margins(cells[0], 90, 110, 90, 110)
        set_cell_margins(cells[1], 90, 140, 90, 140)
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(number.zfill(2))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(WHITE)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(name.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(TEXT)
    set_table_borders(table, WHITE, "8")


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(len(rows), len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_fill(cell, NAVY if row_index == 0 else ICE if row_index % 2 else WHITE)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value)
            for run in p.runs:
                run.font.size = Pt(8.5)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(document: Document, alt: str, target: str) -> None:
    image_path = (SOURCE.parent / target).resolve()
    if not image_path.exists():
        return
    width = Inches(2.65) if "mobile" in image_path.name else Inches(6.6)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    add_picture(p.add_run(), image_path, width=width)
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(9)
    run = caption.add_run(alt)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def render_markdown(document: Document, source: str) -> None:
    lines = source.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    skipped_hero = False
    skip_metadata = True

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.35)
                p.paragraph_format.right_indent = Cm(0.35)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                shading = parse_xml(
                    rf'<w:shd {nsdecls("w")} w:fill="{NAVY}"/>'
                )
                p._p.get_or_add_pPr().append(shading)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(8.3)
                run.font.color.rgb = RGBColor.from_string(WHITE)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if stripped.startswith("# "):
            index += 1
            continue
        if stripped == "## Технічний опис системи":
            index += 1
            continue
        if skip_metadata and (
            stripped.startswith("**Версія документа:**")
            or stripped.startswith("**Стан продукту:**")
            or stripped.startswith("**Дата:**")
        ):
            index += 1
            continue
        skip_metadata = False

        image = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image:
            if "my-meetings-light" in image.group(2) and not skipped_hero:
                skipped_hero = True
            else:
                add_image(document, image.group(1), image.group(2))
            index += 1
            continue

        if stripped.startswith("## "):
            p = document.add_paragraph(style="Heading 1")
            p.paragraph_format.space_before = Pt(16)
            add_inline(p, stripped[3:].upper())
            for run in p.runs:
                force_run_font(run, "Arial")
            set_paragraph_bottom_border(p, CYAN, "14")
            set_repeatable_heading(p)
            index += 1
            continue
        if stripped.startswith("### "):
            p = document.add_paragraph(style="Heading 2")
            add_inline(p, stripped[4:])
            for run in p.runs:
                force_run_font(run, "Arial")
            set_repeatable_heading(p)
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines: list[str] = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_lines.append(candidate)
                index += 1
            rows = [
                [cell.strip() for cell in row.strip("|").split("|")]
                for row in table_lines
                if not re.fullmatch(r"\|?[\s:|-]+\|?", row)
                and not all(
                    re.fullmatch(r":?-{3,}:?", cell.strip())
                    for cell in row.strip("|").split("|")
                )
            ]
            add_markdown_table(document, rows)
            continue

        if stripped.startswith("- "):
            while index < len(lines) and lines[index].strip().startswith("- "):
                p = document.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.65)
                p.paragraph_format.first_line_indent = Cm(-0.25)
                p.paragraph_format.space_after = Pt(3)
                add_inline(p, lines[index].strip()[2:])
                index += 1
            continue

        if re.match(r"\d+\.\s+", stripped):
            while index < len(lines) and re.match(r"\d+\.\s+", lines[index].strip()):
                match = re.match(r"(\d+)\.\s+(.+)", lines[index].strip())
                assert match is not None
                number, text = match.groups()
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.65)
                p.paragraph_format.first_line_indent = Cm(-0.25)
                p.paragraph_format.space_after = Pt(3)
                number_run = p.add_run(f"{number}.  ")
                number_run.bold = True
                number_run.font.color.rgb = RGBColor.from_string(COBALT)
                add_inline(p, text)
                index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if (
                not candidate
                or candidate.startswith("#")
                or candidate.startswith("```")
                or candidate.startswith("- ")
                or re.match(r"\d+\.\s+", candidate)
                or candidate.startswith("|")
                or re.fullmatch(r"!\[(.+?)\]\((.+?)\)", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        p = document.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))


def main() -> None:
    if not SOURCE.exists():
        raise SystemExit(f"Missing source: {SOURCE}")
    if not BRAND_MARK.exists():
        raise SystemExit(
            "Missing docs/screenshots/brand-mark.png; run "
            "`node scripts/capture-doc-screenshots.mjs` first."
        )

    source = SOURCE.read_text(encoding="utf-8")
    sections = [
        line[3:].strip()
        for line in source.splitlines()
        if line.startswith("## ") and line != "## Технічний опис системи"
    ]

    document = Document()
    configure_document(document)
    props = document.core_properties
    props.title = "MeetBroker — технічний опис системи"
    props.subject = "Архітектура, дані, безпека, інтеграції та CI"
    props.author = "MeetBroker contributors"
    props.keywords = "MeetBroker, React, NestJS, PostgreSQL, Docker, booking"

    add_cover(document)
    document.add_page_break()
    add_contents(document, sections)
    document.add_page_break()
    render_markdown(document, source)

    document.save(OUTPUT)
    print(f"Created {OUTPUT.relative_to(ROOT)}")


if __name__ == "__main__":
    try:
        main()
    except Exception as error:
        print(f"DOCX build failed: {error}", file=sys.stderr)
        raise
